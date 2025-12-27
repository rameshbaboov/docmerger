import os
import csv
from docx import Document

# === CONFIGURATION ===
INPUT_FOLDER = "input_docs"         # folder containing .docx files
OUTPUT_FOLDER = "merged_output"     # folder where merged file is saved
OUTPUT_FILE = "merged1.docx"         # name of the merged file
PROCESSED_FILE = "processed1.csv"    # log file (filename, status)


def load_processed_files():
    processed = {}
    if os.path.exists(PROCESSED_FILE):
        with open(PROCESSED_FILE, newline='', encoding="utf-8") as f:
            reader = csv.reader(f)
            for row in reader:
                if len(row) == 2:
                    processed[row[0]] = row[1]
    return processed


def update_processed_file(filename, status):
    with open(PROCESSED_FILE, "a", newline='', encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow([filename, status])


def append_document(master, filename):
    """Append content of a sub-document to master without raw XML insertions."""
    sub_doc = Document(filename)

    # Copy paragraphs
    for para in sub_doc.paragraphs:
        new_para = master.add_paragraph()
        for run in para.runs:
            new_run = new_para.add_run(run.text)
            # Preserve basic formatting
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            if run.font.size:
                new_run.font.size = run.font.size
            if run.font.name:
                new_run.font.name = run.font.name

    # Copy tables
    for table in sub_doc.tables:
        master.add_paragraph()  # spacing before table
        master.element.body.append(table._element)


def merge_docx_files():
    os.makedirs(OUTPUT_FOLDER, exist_ok=True)
    output_path = os.path.join(OUTPUT_FOLDER, OUTPUT_FILE)

    processed = load_processed_files()

    if os.path.exists(output_path):
        merged_document = Document(output_path)
    else:
        merged_document = Document()
        merged_document._element.body.clear_content()

    for filename in sorted(os.listdir(INPUT_FOLDER)):
        if not filename.endswith(".docx"):
            continue
        if filename in processed:
            continue

        file_path = os.path.join(INPUT_FOLDER, filename)

        try:
            if merged_document.paragraphs:  # only if not the very first file
                merged_document.add_page_break()

            append_document(merged_document, file_path)

            merged_document.save(output_path)
            update_processed_file(filename, "success")
            print(f"Processed: {filename}")
        except Exception as e:
            update_processed_file(filename, "error")
            print(f"Error processing {filename}: {e}")


if __name__ == "__main__":
    merge_docx_files()
