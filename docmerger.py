import time
import os
import csv
import argparse
import logging
from pathlib import Path
from docx import Document
from docx.enum.text import WD_BREAK

# === CONFIGURATION (defaults) ===
INPUT_FOLDER = "input_docs"          # folder containing .docx files
OUTPUT_FOLDER = "merged_output"      # folder where merged file is saved
OUTPUT_FILE = "merged.docx"          # name of the merged file
PROCESSED_FILE = "processed.csv"     # log file (filename, status)


logger = logging.getLogger("docmerger")


def _setup_logging(output_folder: str) -> None:
    """Configure file logging to merged_output/docmerger.log."""
    Path(output_folder).mkdir(parents=True, exist_ok=True)
    log_path = Path(output_folder) / "docmerger.log"

    # Avoid adding handlers multiple times if imported.
    if logger.handlers:
        return

    logger.setLevel(logging.INFO)
    fh = logging.FileHandler(log_path, encoding="utf-8")
    fmt = logging.Formatter("%(asctime)s %(levelname)s %(message)s")
    fh.setFormatter(fmt)
    logger.addHandler(fh)


def load_processed_files():
    """Read processed.csv and return dict of {filename: status}."""
    processed = {}
    if os.path.exists(PROCESSED_FILE):
        with open(PROCESSED_FILE, newline='', encoding="utf-8") as f:
            reader = csv.reader(f)
            for row in reader:
                if len(row) == 2:
                    processed[row[0]] = row[1]
    return processed


def update_processed_file(filename, status):
    """Append entry to processed.csv."""
    with open(PROCESSED_FILE, "a", newline='', encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow([filename, status])


def merge_docx_files():
    os.makedirs(INPUT_FOLDER, exist_ok=True)
    os.makedirs(OUTPUT_FOLDER, exist_ok=True)
    _setup_logging(OUTPUT_FOLDER)

    logger.info(
        "merge pass started input_folder=%s output_folder=%s output_file=%s processed_file=%s",
        INPUT_FOLDER,
        OUTPUT_FOLDER,
        OUTPUT_FILE,
        PROCESSED_FILE,
    )
    output_path = os.path.join(OUTPUT_FOLDER, OUTPUT_FILE)

    # Load processed status
    processed = load_processed_files()

    # Load or create merged document
    if os.path.exists(output_path):
        merged_document = Document(output_path)
    else:
        merged_document = Document()

    # Process new files
    for filename in sorted(os.listdir(INPUT_FOLDER)):
        if not filename.endswith(".docx"):
            continue

        if filename in processed:
            # Skip already processed files (success or error)
            continue

        file_path = os.path.join(INPUT_FOLDER, filename)

        try:
            sub_doc = Document(file_path)

            # Add page break if not first content in merged doc
            if merged_document.paragraphs:
                merged_document.add_page_break()

            # Append each element from sub_doc
            for element in sub_doc.element.body:
                merged_document.element.body.append(element)

            merged_document.save(output_path)
            update_processed_file(filename, "success")
            logger.info("processed filename=%s status=success", filename)
            print(f"Processed: {filename}")

        except Exception as e:
            update_processed_file(filename, "error")
            logger.exception("processed filename=%s status=error error=%s", filename, str(e))
            print(f"Error processing {filename}: {e}")

    logger.info("merge pass finished")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Merge .docx files from a folder into one output .docx")
    parser.add_argument("--input-folder", default=INPUT_FOLDER)
    parser.add_argument("--output-folder", default=OUTPUT_FOLDER)
    parser.add_argument("--output-file", default=OUTPUT_FILE)
    parser.add_argument("--processed-file", default=PROCESSED_FILE)
    parser.add_argument("--interval", type=int, default=300, help="Seconds between merge passes")
    parser.add_argument("--once", action="store_true", help="Run a single merge pass and exit")
    args = parser.parse_args()

    # Apply CLI overrides (keeps backwards compatibility for code that imports this module)
    INPUT_FOLDER = args.input_folder
    OUTPUT_FOLDER = args.output_folder
    OUTPUT_FILE = args.output_file
    PROCESSED_FILE = args.processed_file

    if args.once:
        merge_docx_files()
    else:
        while True:
            merge_docx_files()
            logger.info("sleeping seconds=%s", args.interval)
            print(f"Waiting {args.interval} seconds before next run...")
            time.sleep(args.interval)
